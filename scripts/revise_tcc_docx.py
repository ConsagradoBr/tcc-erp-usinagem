from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE_DOC = Path(r"D:\Download\TCC-Word.docx")
DEFAULT_OUTPUT_DOC = ROOT / "output" / "doc" / "TCC-Word-alinhado.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def remove_all_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_run_font(
    run,
    name: str,
    size_pt: int,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def normalize_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading.font.size = Pt(14)
    heading.font.bold = True

    for style_name in ["toc 1", "toc 2", "toc 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)


def normalize_body_paragraph(paragraph, index: int) -> None:
    text = paragraph.text
    style_name = paragraph.style.name

    if index >= 39 and style_name in {"Normal", "Normal (Web)"} and text.strip():
        paragraph.style = "Normal"
        remove_all_runs(paragraph)
        run = paragraph.add_run(text)
        set_run_font(run, "Times New Roman", 12)
        return

    if style_name == "Heading 1" and text.strip():
        remove_all_runs(paragraph)
        run = paragraph.add_run(text)
        set_run_font(run, "Times New Roman", 14, bold=True)
        return

    if style_name.startswith("toc ") and text.strip():
        remove_all_runs(paragraph)
        run = paragraph.add_run(text)
        set_run_font(
            run,
            "Times New Roman",
            10,
            bold=(style_name == "toc 1"),
            italic=(style_name == "toc 3"),
        )


def cleanup_artifact_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Pronto") and "GitHub mostra hoje" in text:
            delete_paragraph(paragraph)

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        delete_paragraph(doc.paragraphs[-1])


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Revisa e alinha o .docx do TCC.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE_DOC)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DOC)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source_doc = args.source
    output_doc = args.output

    output_doc.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(source_doc)
    paragraphs = doc.paragraphs

    replacements = {
        41: (
            "Este Trabalho de Conclusao de Curso apresenta o desenvolvimento de um sistema ERP para a AMP Usinagem, "
            "voltado a pequenas e medias empresas do setor de usinagem. O projeto surgiu da necessidade de reduzir "
            "retrabalho em cadastros, organizar o fluxo comercial e produtivo e ampliar o controle financeiro e fiscal. "
            "A solucao foi estruturada com React 18 e Vite no frontend, Flask no backend e PostgreSQL como base da "
            "arquitetura web, mantendo suporte a SQLite para execucao local em modo desktop. Entre as funcionalidades "
            "entregues estao autenticacao segura, gestao de clientes com importacao de NF-e, orcamentos, ordens de "
            "servico, controle financeiro com parcelamento e leitura de boletos PDF, usuarios, backup e dashboard "
            "analitico. O resultado foi um produto funcional, com distribuicao desktop nativa para Windows, capaz de "
            "centralizar processos e aumentar a rastreabilidade operacional da empresa."
        ),
        46: (
            "This Course Conclusion Work presents the development of an ERP system for AMP Usinagem, aimed at small and "
            "medium-sized machining companies. The project was created to reduce rework in records, organize commercial "
            "and production flows, and improve financial and fiscal control. The solution was built with React 18 and "
            "Vite on the frontend, Flask on the backend, and PostgreSQL as the main database architecture, while also "
            "supporting SQLite for local desktop execution. Delivered features include secure authentication, customer "
            "management with NF-e import, quotations, service orders, financial control with installments and PDF bank "
            "slip parsing, user management, backup, and an analytical dashboard. The final result is a functional "
            "product with native desktop distribution for Windows, capable of centralizing processes and increasing "
            "operational traceability in the partner company."
        ),
        70: (
            "A empresa de usinagem AMP Usinagem, especializada na producao de pecas mecanicas de alta precisao, opera em "
            "um ambiente onde a organizacao dos processos influencia diretamente prazo, qualidade e margem. Como em muitas "
            "PMEs do setor, parte do controle historicamente foi feita com planilhas, mensagens, documentos fiscais e "
            "rotinas manuais, o que gera retrabalho, informacao dispersa e baixa rastreabilidade entre comercial, "
            "producao e financeiro."
        ),
        71: (
            "Diante desse contexto, este trabalho tem como objetivo desenvolver um sistema de gestao integrado para a AMP "
            "Usinagem, reunindo autenticacao, clientes, orcamentos, ordens de servico, financeiro e indicadores em uma "
            "mesma plataforma. A proposta privilegia a operacao real da empresa, com importacao de dados fiscais, fluxo "
            "de aprovacao comercial e visao consolidada da operacao."
        ),
        72: (
            "A justificativa do projeto esta na necessidade de transformar um processo fragmentado em uma experiencia "
            "integrada, com menos erros operacionais, menor dependencia de digitacao manual e melhor capacidade de "
            "decisao. O estudo delimita-se ao desenvolvimento e validacao do nucleo do ERP, mantendo como evolucoes "
            "futuras a ampliacao do dominio de fornecedores, estoque e integracoes externas adicionais."
        ),
        74: (
            "O tema deste trabalho e o desenvolvimento de um sistema de gestao para uma empresa de usinagem, com foco na "
            "integracao entre relacionamento com clientes, orcamentos, ordens de servico, controle financeiro, leitura "
            "de documentos fiscais e visualizacao de indicadores. A solucao foi construida com Python e Flask no backend, "
            "React no frontend e arquitetura de dados baseada em PostgreSQL, mantendo suporte a execucao desktop local "
            "para demonstracao e operacao."
        ),
        76: (
            "Empresas de usinagem frequentemente gerenciam pedidos, clientes, producao e financeiros com ferramentas "
            "desconexas, como planilhas, e-mails e documentos fiscais. Esse modelo dificulta o rastreamento do pedido "
            "desde o contato comercial ate o recebimento, aumenta a chance de erros de cadastro e compromete a visao "
            "integrada da operacao. O problema central deste trabalho e como transformar esses fluxos dispersos em um "
            "sistema unico, simples de usar e aderente a rotina real da empresa."
        ),
        78: (
            "A relevancia do projeto esta na necessidade de digitalizacao do setor de usinagem com uma abordagem acessivel "
            "e orientada a resultado. O sistema proposto reduz tarefas manuais, melhora a rastreabilidade fiscal e "
            "financeira e cria uma base mais segura para planejamento e acompanhamento da producao. Para o TCC, o projeto "
            "demonstra a aplicacao pratica de conceitos de engenharia de software, UX e integracao de processos de negocio."
        ),
        80: (
            "Desenvolver um sistema de gestao para a AMP Usinagem, integrando autenticacao, clientes, orcamentos, ordens "
            "de servico, financeiro e dashboard analitico, com foco em eficiencia operacional, rastreabilidade e "
            "distribuicao desktop."
        ),
        83: "Projetar e implementar o backend em Python e Flask para sustentar APIs de autenticacao, clientes, orcamentos, ordens de servico, financeiro, usuarios e backup.",
        84: "Desenvolver o frontend com React para oferecer uma interface clara, responsiva e alinhada ao fluxo real de operacao da empresa.",
        85: "Estruturar a persistencia de dados com PostgreSQL como arquitetura principal e suporte a SQLite em execucao local desktop.",
        86: "Implementar importacoes nativas de NF-e e leitura de boletos PDF diretamente no backend, sem dependencia de ferramentas externas como N8N.",
        87: "Validar o sistema em cenarios proximos da operacao real e documentar a evolucao entre prototipo, implementacao e resultados.",
        89: (
            "Como apoio a validacao do projeto, foram produzidos prototipos, capturas das telas e demonstracoes praticas "
            "do sistema para comparar o fluxo inicial idealizado com a implementacao real entregue."
        ),
        90: (
            "Esses materiais servem como evidencia do amadurecimento do produto e apoiam a apresentacao do trabalho para "
            "a banca e para a empresa parceira."
        ),
        103: (
            "PostgreSQL hospedado no Supabase como arquitetura principal da versao web, garantindo integridade referencial, "
            "transacoes ACID e escalabilidade. Para distribuicao desktop e operacao local, o sistema tambem suporta SQLite "
            "como persistencia fallback."
        ),
        108: (
            "O desenvolvimento do sistema foi realizado por uma equipe de tres integrantes ao longo de aproximadamente doze "
            "meses. Para organizar as entregas, foi adotada uma rotina inspirada em Scrum, com ciclos iterativos de "
            "levantamento de requisitos, prototipacao, implementacao e validacao."
        ),
        110: (
            "Em seguida, foi realizada a etapa de prototipacao no Figma, onde foram criadas e validadas as telas principais "
            "de login, dashboard, clientes e financeiro, incluindo variacoes light e dark. Esses prototipos serviram como "
            "base de identidade visual e orientaram a simplificacao da experiencia do produto."
        ),
        111: (
            "Na fase de desenvolvimento, o sistema foi dividido em frontend e backend. O frontend foi construido com React "
            "e organizado em modulos operacionais. O backend foi desenvolvido com Flask para concentrar regras de negocio, "
            "autenticacao, importacoes e persistencia. Ao final, o projeto tambem recebeu empacotamento desktop nativo para Windows."
        ),
        112: (
            "O banco de dados principal considerado na arquitetura foi o PostgreSQL, sem abrir mao da possibilidade de "
            "execucao local com SQLite no aplicativo desktop."
        ),
        119: (
            "Realizamos entrevistas com responsaveis da AMP Usinagem e analise dos processos manuais atuais. Os requisitos "
            "funcionais priorizados envolveram autenticacao, cadastro de clientes, importacao de NF-e, orcamentos, fluxo "
            "de ordens de servico, lancamentos financeiros com parcelamento, usuarios e backup."
        ),
        121: (
            "Com base nos requisitos, criamos prototipos de alta fidelidade no Figma para validar login, dashboard, "
            "financeiro e navegacao lateral. Os prototipos ajudaram a consolidar a identidade AMP e serviram como base "
            "para refinamentos posteriores do sistema."
        ),
        124: (
            "Frontend: React 18 com Vite como ferramenta de build e desenvolvimento, Tailwind CSS v3 para estilizacao, "
            "React Router DOM v6 para navegacao SPA, Axios para chamadas a API, Recharts para indicadores e react-hot-toast "
            "para feedbacks de interface."
        ),
        126: (
            "Backend: Python com Flask 3.1.0, integrado a Flask-SQLAlchemy 3.1.1, Flask-JWT-Extended 4.7.0, Flask-Cors, "
            "psycopg2-binary e pdfplumber. Essa base sustenta autenticacao, regras de negocio, importacoes e integracoes "
            "internas do sistema."
        ),
        128: (
            "Banco de Dados: PostgreSQL hospedado no Supabase como base principal da arquitetura web, com SQLite como "
            "fallback para a distribuicao local no aplicativo desktop."
        ),
        140: (
            "O sistema foi implementado com sucesso, entregando uma solucao ERP funcional para a AMP Usinagem. A arquitetura "
            "separada entre frontend e backend, somada ao empacotamento desktop, permitiu evolucao modular e aderencia ao "
            "uso real da empresa."
        ),
        142: (
            "Autenticacao, usuarios e acesso: login com JWT, rotas protegidas, bootstrap do primeiro administrador e "
            "gerenciamento interno de usuarios com perfis de acesso."
        ),
        146: (
            "Gestao de clientes: CRUD completo, busca e filtro, importacao automatica de NF-e (.xml ou .json), deteccao "
            "de duplicidade por documento e consolidacao do relacionamento comercial."
        ),
        150: (
            "Controle financeiro: lancamentos a pagar e a receber com parcelamento, baixa individual ou em grupo, "
            "importacao de boletos PDF, calculo de juros, identificacao de status e apoio ao fluxo comercial."
        ),
        154: (
            "Orcamentos, ordens de servico e dashboard: o sistema conecta aprovacao comercial, andamento produtivo e "
            "visao financeira em uma mesma plataforma, com indicadores em tempo real, grafico de receitas x pagamentos e "
            "acompanhamento de situacoes operacionais."
        ),
        158: (
            "O sistema automatiza tarefas criticas da rotina da empresa: reduz tempo de cadastro por importacao de NF-e, "
            "elimina grande parte da digitacao manual no financeiro e organiza a transicao entre orcamento, ordem de servico "
            "e recebimento."
        ),
        159: (
            "Testes unitarios, de integracao e funcionais foram executados sobre fluxos essenciais, como autenticacao, "
            "parcelamento, importacoes e atualizacao de status, confirmando a estabilidade do nucleo entregue."
        ),
        161: (
            "Comparado ao processo manual anterior, o ERP centraliza dados, melhora a rastreabilidade fiscal e financeira "
            "e reduz o atrito operacional. Como evolucoes futuras, destacam-se a integracao completa de fornecedores ao "
            "dominio de relacionamento, estoque, integracao SEFAZ, relatorios produtivos avancados e expansao mobile."
        ),
        165: (
            "O objetivo principal do trabalho foi alcancado: desenvolver um sistema ERP funcional para a AMP Usinagem, "
            "capaz de integrar autenticacao, clientes, orcamentos, ordens de servico, financeiro e indicadores em uma "
            "mesma plataforma."
        ),
        166: "Foram implementados com sucesso:",
        167: "Autenticacao segura via JWT, com controle de acesso e gestao interna de usuarios;",
        168: "Gestao de clientes com importacao automatica de NF-e e deteccao de duplicidades;",
        169: "Controle financeiro com lancamentos parcelados, leitura de boletos PDF, juros, status e vinculos com o fluxo comercial;",
        170: "Dashboard analitico, orcamentos, ordens de servico, backup e distribuicao desktop nativa para Windows.",
        171: (
            "A stack escolhida - React 18 + Vite no frontend, Flask no backend e PostgreSQL como base principal da "
            "arquitetura - garantiu uma aplicacao responsiva, segura e extensivel, com suporte adicional a execucao local "
            "em SQLite no ambiente desktop."
        ),
        172: (
            "Os testes confirmaram o funcionamento dos fluxos principais e a aderencia aos requisitos priorizados ao longo "
            "do projeto."
        ),
        173: (
            "Como continuidade, o projeto pode evoluir com integracao completa de fornecedores, estoque, notificacoes em "
            "tempo real, integracao SEFAZ, relatorios OEE e expansao mobile."
        ),
        174: (
            "O sistema entregue representa uma solucao pratica, utilizavel e escalavel para a empresa parceira, ao mesmo "
            "tempo em que demonstra o amadurecimento do prototipo inicial para um produto com aplicacao real."
        ),
    }

    manual_toc = {
        49: "Sumario",
        50: "Capa .................................................................. 1",
        51: "1 INTRODUCAO .......................................................... 6",
        52: "1.1 Tema .............................................................. 6",
        53: "1.2 Problema .......................................................... 7",
        54: "1.3 Justificativa ..................................................... 7",
        55: "1.3.1 Acoes ........................................................... 8",
        56: "2 REFERENCIAL TEORICO ................................................. 9",
        57: "2.1 Tecnologias utilizadas no desenvolvimento ......................... 9",
        58: "2.1.1 HTML ............................................................ 9",
        59: "2.1.2 CSS ............................................................. 9",
        60: "2.1.3 React (Frontend) ................................................ 9",
        61: "2.1.4 Flask (Backend) ................................................  9",
        62: "2.1.5 Banco de Dados ................................................. 10",
        63: "2.1.6 Automacoes ..................................................... 10",
        64: "3 METODOLOGIA ........................................................ 11",
        65: "4 RESULTADOS E DISCUSSAO ............................................. 12",
        66: "5 CONCLUSAO .......................................................... 13",
        67: "",
    }

    for index, text in replacements.items():
        replace_paragraph_text(paragraphs[index], text)

    for index, text in manual_toc.items():
        paragraph = paragraphs[index]
        paragraph.style = "Normal"
        clear_paragraph_content(paragraph)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, "Times New Roman", 12)

    normalize_document_styles(doc)
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph is None:
            continue
        normalize_body_paragraph(paragraph, index)

    cleanup_artifact_paragraphs(doc)

    doc.save(output_doc)
    print(f"Saved revised document to: {output_doc}")


if __name__ == "__main__":
    main()
